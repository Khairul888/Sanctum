from docx import Document


def load_docx(file_path):
    document = Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)
